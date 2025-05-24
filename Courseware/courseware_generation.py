"""
File: courseware_generation.py

===============================================================================
Courseware Document Generator
===============================================================================
Description:
    This module serves as the main entry point for the Courseware Document Generator
    application. It is designed to parse Course Proposal (CP) documents, extract and
    interpret the course data, and generate multiple courseware documents such as:
      - Learning Guide (LG)
      - Assessment Plan (AP)
      - Lesson Plan (LP)
      - Facilitator's Guide (FG)
      - Timetable (as needed)
      
    The application utilizes both AI-based processing (via OpenAI and autogen agents)
    and conventional document parsing and web scraping methods to ensure that the CP data
    is accurately transformed into a structured format for document generation.

Main Functionalities:
    1. Data Models:
        - Defines several Pydantic models (e.g., Topic, LearningUnit, CourseData, etc.)
          to validate and structure the course proposal and generated document data.
          
    2. Document Parsing:
        - Function: parse_cp_document(uploaded_file)
          Parses a CP document (Word or Excel) into a trimmed Markdown string based on
          regex patterns to capture only the relevant sections of the document.
          
    3. Web Scraping:
        - Function: web_scrape(course_title, name_of_org)
          Automates a headless browser session using Selenium to retrieve TGS Ref No (and UEN)
          from the MySkillsFuture portal based on the provided course title and organization.
          
    4. Data Interpretation:
        - Function: interpret_cp(raw_data, model_client)
          Leverages an AI assistant (via the OpenAIChatCompletionClient) to extract and structure
          the course proposal data into a comprehensive JSON dictionary as defined by the CourseData model.
          
    5. Streamlit Application:
        - Function: app()
          Implements the user interface using Streamlit. This interface guides users through:
            - Uploading a Course Proposal document.
            - Managing organization details (CRUD operations via a modal).
            - Optionally uploading an updated Skills Framework dataset.
            - Selecting which courseware documents to generate.
            - Executing the parsing, data extraction, document generation processes,
              and finally providing a ZIP file download of all generated documents.
              
Dependencies:
    - Custom Courseware Utilities:
        • Courseware.utils.agentic_LG         : For generating the Learning Guide.
        • Courseware.utils.agentic_AP         : For generating Assessment Documents.
        • Courseware.utils.timetable_generator : For generating the course timetable.
        • Courseware.utils.agentic_LP         : For generating the Lesson Plan.
        • Courseware.utils.agentic_FG         : For generating the Facilitator's Guide.
        • Courseware.utils.model_configs       : For model configuration and selection.
        • Courseware.utils.organization_utils  : For managing organization data (CRUD).
    - External Libraries:
        • os, io, zipfile, tempfile, json, time, asyncio, datetime
        • streamlit                        : For building the web UI.
        • selenium & BeautifulSoup         : For web scraping tasks.
        • docx                             : For generating and modifying Word documents.
        • pydantic                         : For data validation and structured models.
        • autogen_agentchat & autogen_core   : For AI-assisted text generation and processing.
        • urllib.parse                     : For URL manipulation.
    
Usage:
    - Configure API keys and endpoints in st.secrets (e.g., LLAMA_CLOUD_API_KEY, BROWSER_TOKEN,
      BROWSER_WEBDRIVER_ENDPOINT, etc.).
    - Run this module using Streamlit, e.g., `streamlit run <this_file.py>`, to launch the web interface.
    - Follow the on-screen instructions to upload your CP document, manage organization data, select
      the desired courseware documents, and generate/download the outputs.

Author: 
    Derrick Lim
Date:
    4 March 2025

Notes:
    - This module uses asynchronous functions and external AI services for data extraction.
    - The Selenium web scraping component is configured to run headlessly with optimized options
      suitable for both local and containerized environments.
    - Organization management is performed using a JSON-based system via utility functions provided
      in the Courseware.utils.organization_utils module.
    - Ensure all dependencies are installed and properly configured before running the application.

===============================================================================
"""


from Courseware.utils.agentic_LG import generate_learning_guide
from Courseware.utils.agentic_AP import generate_assessment_documents
from Courseware.utils.timetable_generator import generate_timetable
from Courseware.utils.agentic_LP import generate_lesson_plan
from Courseware.utils.agentic_FG import generate_facilitators_guide
from Courseware.utils.model_configs import MODEL_CHOICES, get_model_config
import os
import io
import zipfile
import tempfile
import json 
import time
import asyncio
from datetime import datetime
import streamlit as st
import urllib.parse
from selenium import webdriver
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from selenium import webdriver
from bs4 import BeautifulSoup
from pydantic import BaseModel
from typing import List, Optional
from autogen_agentchat.agents import AssistantAgent
from autogen_agentchat.messages import TextMessage
from autogen_core import CancellationToken
from autogen_ext.models.openai import OpenAIChatCompletionClient
from utils.helper import save_uploaded_file, parse_json_content
# Import organisation CRUD utilities and model
from Courseware.utils.organization_utils import (
    load_organizations,
    save_organizations,
    add_organization,
    update_organization,
    delete_organization,
    Organization
)
from streamlit_modal import Modal

# Initialize session state variables
if 'lg_output' not in st.session_state:
    st.session_state['lg_output'] = None
if 'ap_output' not in st.session_state:
    st.session_state['ap_output'] = None
if 'lp_output' not in st.session_state:
    st.session_state['lp_output'] = None
if 'fg_output' not in st.session_state:
    st.session_state['fg_output'] = None
if 'context' not in st.session_state:
    st.session_state['context'] = None
if 'asr_output' not in st.session_state:
    st.session_state['asr_output'] = None
if 'selected_model' not in st.session_state:
    st.session_state['selected_model'] = "Gemini-Pro-2.5-Exp-03-25"

############################################################
# 1. Pydantic Models
############################################################
class Topic(BaseModel):
    Topic_Title: str
    Bullet_Points: List[str]

class KDescription(BaseModel):
    K_number: str
    Description: str

class ADescription(BaseModel):
    A_number: str
    Description: str

class LearningUnit(BaseModel):
    LU_Title: str
    LU_Duration: str  # <-- Add this field
    Topics: List[Topic]
    LO: str
    K_numbering_description: List[KDescription]
    A_numbering_description: List[ADescription]
    Assessment_Methods: List[str]
    Instructional_Methods: List[str]

class EvidenceDetail(BaseModel):
    LO: str
    Evidence: str

class AssessmentMethodDetail(BaseModel):
    Assessment_Method: str
    Method_Abbreviation: str
    Total_Delivery_Hours: str
    Assessor_to_Candidate_Ratio: List[str]
    Evidence: Optional[List[EvidenceDetail]] = None
    Submission: Optional[List[str]] = None
    Marking_Process: Optional[List[str]] = None
    Retention_Period: Optional[str] = None

class CourseData(BaseModel):
    Date: str 
    Year: str
    Name_of_Organisation: str
    Course_Title: str
    TSC_Title: str
    TSC_Code: str
    Total_Training_Hours: str 
    Total_Assessment_Hours: str 
    Total_Course_Duration_Hours: str 
    Learning_Units: List[LearningUnit]
    Assessment_Methods_Details: List[AssessmentMethodDetail]

class Session(BaseModel):
    Time: str
    instruction_title: str
    bullet_points: List[str]
    Instructional_Methods: str
    Resources: str

class DayLessonPlan(BaseModel):
    Day: str
    Sessions: List[Session]

class LessonPlan(BaseModel):
    lesson_plan: List[DayLessonPlan]

############################################################
# 2. Course Proposal Document Parsing
############################################################
from llama_cloud_services import LlamaParse
from llama_index.core import SimpleDirectoryReader
import os
import re

# Flatten Instructional Methods Column which has new lines
def flatten_instructional_methods_in_curriculum_table(docx_path, output_path=None):
    doc = Document(docx_path)
    for table in doc.tables:
        # Get header cells' text for this table
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        # Check if this is the curriculum key features table by matching key headers
        if (
            "S/N" in header_cells and
            "LUs" in header_cells and
            "LOs*" in header_cells and
            "Instructional Methods" in header_cells
        ):
            # Only flatten this table's "Instructional Methods" column
            idx = header_cells.index("Instructional Methods")
            for row in table.rows[1:]:
                cell = row.cells[idx]
                # Flatten by replacing newlines with commas
                cell.text = ", ".join(line.strip() for line in cell.text.splitlines() if line.strip())
            break  # Stop after processing the first matching table
    if output_path:
        doc.save(output_path)
    else:
        doc.save(docx_path)
        
    # Flatten post production
    import re

def flatten_markdown_table_column(md_text, column_name="Instructional Methods"):
    lines = md_text.splitlines()
    new_lines = []
    in_curriculum_table = False
    header = None
    col_idx = None

    for i, line in enumerate(lines):
        # Detect the curriculum table header
        if (
            "|S/N|" in line and
            "|LUs|" in line and
            "|LOs*" in line and
            f"|{column_name}|" in line
        ):
            in_curriculum_table = True
            header = [h.strip() for h in line.split("|")]
            col_idx = header.index(column_name)
            new_lines.append(line)
            continue

        # Detect the end of the table (next non-table line or empty line)
        if in_curriculum_table and (not line.strip().startswith("|") or line.strip() == ""):
            in_curriculum_table = False

        # If inside the curriculum table, flatten the column
        if in_curriculum_table and "|" in line and not line.strip().startswith("|---"):
            cells = [c.strip() for c in line.split("|")]
            if col_idx is not None and len(cells) > col_idx:
                # Flatten the cell (replace newlines and extra spaces with commas)
                cells[col_idx] = re.sub(r"\s*\n\s*", ", ", cells[col_idx]).replace("  ", " ")
                new_line = "|".join(cells)
                new_lines.append(new_line)
            else:
                new_lines.append(line)
        else:
            new_lines.append(line)

    return "\n".join(new_lines)

def optimize_timetable(context):
    """
    Optimizes the timetable to ensure each Learning Unit fully utilizes its allocated duration.
    
    Args:
        context (dict): The course context containing Learning Units and lesson plan
        
    Returns:
        dict: Updated context with optimized lesson plan
    """
    import re
    from copy import deepcopy
    from datetime import datetime, timedelta
    
    # Make a deep copy to avoid modifying the original data
    context = deepcopy(context)
    learning_units = context['Learning_Units']
    lesson_plan = context['lesson_plan']
    
    # Helper function to extract start time for sorting
    def extract_time_for_sorting(time_str):
        match = re.search(r'(\d{4})hrs', time_str)
        if match:
            return match.group(1)
        return "0000"

    # Extract LU durations in minutes
    lu_durations = {}
    for lu in learning_units:
        match = re.search(r'(\d+(\.\d+)?)', lu['LU_Duration'])
        if match:
            hours = float(match.group(1))
            lu_durations[lu['LU_Title']] = int(hours * 60)
    
    # Function to extract duration from a time string (e.g., "0935hrs - 1100hrs (1h 25min)")
    def extract_duration_minutes(time_str):
        # Handle different formats
        duration_match = re.search(r'\((\d+)\s*min', time_str)
        if duration_match:
            return int(duration_match.group(1))
            
        # Look for format like (1h 25min)
        duration_match = re.search(r'\((\d+)h\s*(\d+)\s*min', time_str)
        if duration_match:
            hours = int(duration_match.group(1))
            minutes = int(duration_match.group(2))
            return hours * 60 + minutes
            
        # Look for format like (5 mins)
        duration_match = re.search(r'\((\d+)\s*mins', time_str)
        if duration_match:
            return int(duration_match.group(1))
            
        # If all else fails, extract from time range
        time_range_match = re.search(r'(\d{4})hrs\s*-\s*(\d{4})hrs', time_str)
        if time_range_match:
            start_time = time_range_match.group(1)
            end_time = time_range_match.group(2)
            
            start_hours, start_minutes = parse_time(start_time)
            end_hours, end_minutes = parse_time(end_time)
            
            # Handle case where end time is on the next day
            if end_hours < start_hours:
                end_hours += 24
                
            total_minutes = (end_hours * 60 + end_minutes) - (start_hours * 60 + start_minutes)
            return total_minutes
            
        return 0
    
    # Function to parse time from a string like "0935hrs" or "00hrs"
    def parse_time(time_str):
        # Remove 'hrs' suffix
        cleaned_str = time_str.replace('hrs', '')
        
        # Handle case where time is just hours with no minutes (e.g., "00hrs")
        if len(cleaned_str) <= 2:
            return int(cleaned_str), 0
        
        # Normal case with hours and minutes (e.g., "0935hrs")
        hours = int(cleaned_str[:2])
        minutes = int(cleaned_str[2:])
        return hours, minutes
    
    # Function to format time as "0935hrs"
    def format_time(hours, minutes):
        return f"{hours:02d}{minutes:02d}hrs"
    
    # Function to format duration as "1h 25min" or "30min"
    def format_duration(minutes):
        hours = minutes // 60
        mins = minutes % 60
        if hours > 0 and mins > 0:
            return f"{hours}h {mins} min"
        elif hours > 0:
            return f"{hours}h"
        else:
            return f"{mins} min"
    
    # Function to check if a time is after 1830
    def is_after_end_time(time_str):
        hours, minutes = parse_time(time_str)
        return (hours > 18) or (hours == 18 and minutes > 30)
    
    # Function to find gaps in the schedule (e.g., time before lunch)
    def find_schedule_gaps(day_sessions):
        gaps = []
        sorted_sessions = sorted(day_sessions, key=lambda s: extract_time_for_sorting(s['Time']))
        
        for i in range(len(sorted_sessions) - 1):
            current_session = sorted_sessions[i]
            next_session = sorted_sessions[i+1]
            
            # Skip if current session is administrative
            if current_session.get('is_admin', False):
                continue
                
            # Get end time of current session and start time of next session
            current_end_match = re.search(r'- (\d{4})hrs', current_session['Time'])
            next_start_match = re.search(r'(\d{4})hrs', next_session['Time'])
            
            if current_end_match and next_start_match:
                current_end_time = current_end_match.group(1)
                next_start_time = next_start_match.group(1)
                
                current_end_hours, current_end_minutes = parse_time(current_end_time)
                next_start_hours, next_start_minutes = parse_time(next_start_time)
                
                # Calculate gap in minutes
                current_end_total_minutes = current_end_hours * 60 + current_end_minutes
                next_start_total_minutes = next_start_hours * 60 + next_start_minutes
                gap_minutes = next_start_total_minutes - current_end_total_minutes
                
                # If gap is significant (> 10 min) and not crossing days
                if gap_minutes > 10 and gap_minutes < 120:  # Less than 2 hours to avoid overnight gaps
                    gaps.append({
                        'start_time': current_end_time,
                        'end_time': next_start_time,
                        'duration': gap_minutes,
                        'before_admin': next_session.get('is_admin', False),
                        'index': i+1  # Index where a new session could be inserted
                    })
        
        return gaps
    
    # Identify administrative sessions and final assessments/surveys
    admin_sessions = []
    final_sessions = []  # For assessments and TRAQOM surveys
    
    for day_idx, day in enumerate(lesson_plan):
        for session_idx, session in enumerate(day['Sessions']):
            # Check if this is a break, lunch, or attendance
            if any(keyword in session['instruction_title'].lower() for keyword in 
                   ['break', 'lunch', 'attendance']):
                admin_sessions.append((day_idx, session_idx))
                session['is_admin'] = True
                session['is_final'] = False
            # Check if this is a final assessment or survey
            elif any(keyword in session['instruction_title'].lower() for keyword in 
                    ['assessment', 'traqom survey', 'feedback']):
                final_sessions.append((day_idx, session_idx))
                session['is_admin'] = False
                session['is_final'] = True
            # Regular session
            else:
                session['is_admin'] = False
                session['is_final'] = False
            
            # Initialize LU assignment
            session['lu_assigned'] = None
    
    # Identify which session belongs to which LU
    for day in lesson_plan:
        for session in day['Sessions']:
            # Skip administrative and final sessions
            if session.get('is_admin', False) or session.get('is_final', False):
                continue
            
            # Try to match by topic title
            for lu in learning_units:
                for topic in lu['Topics']:
                    if topic['Topic_Title'] in session['instruction_title']:
                        session['lu_assigned'] = lu['LU_Title']
                        break
                if session.get('lu_assigned'):
                    break
            
            # If not matched by topic, try to identify by LU number in activities
            if session['lu_assigned'] is None and 'Activity' in session['instruction_title']:
                # Look for LU1, LU2, etc. in the title
                for lu in learning_units:
                    lu_number = lu['LU_Title'].split(':')[0].strip()  # e.g., "LU1"
                    if lu_number in session['instruction_title']:
                        session['lu_assigned'] = lu['LU_Title']
                        break
    
    # For activities without explicit LU assignment, infer from context
    for day_idx, day in enumerate(lesson_plan):
        for session_idx, session in enumerate(day['Sessions']):
            if session.get('is_admin', False) or session.get('is_final', False) or session.get('lu_assigned'):
                continue
                
            if 'Activity' in session['instruction_title']:
                # Look backward for context
                for i in range(session_idx - 1, -1, -1):
                    prev_session = day['Sessions'][i]
                    if prev_session.get('lu_assigned'):
                        session['lu_assigned'] = prev_session['lu_assigned']
                        break
                
                # If still not assigned and there's a previous day, check the last topic there
                if session['lu_assigned'] is None and day_idx > 0:
                    prev_day = lesson_plan[day_idx - 1]
                    for prev_session in reversed(prev_day['Sessions']):
                        if prev_session.get('lu_assigned'):
                            session['lu_assigned'] = prev_session['lu_assigned']
                            break
    
    # Manually assign any remaining unassigned content sessions based on context
    for day_idx, day in enumerate(lesson_plan):
        for session_idx, session in enumerate(day['Sessions']):
            if session.get('is_admin', False) or session.get('is_final', False) or session.get('lu_assigned'):
                continue
                
            # Day 1 morning content is likely related to LU1 if not otherwise assigned
            if day_idx == 0:
                lunch_idx = next((i for i, s in enumerate(day['Sessions']) if 'lunch' in s['instruction_title'].lower()), -1)
                if lunch_idx > -1 and session_idx < lunch_idx:
                    session['lu_assigned'] = 'LU1: Catalysing HR with Generative AI (GAI)'
                else:
                    session['lu_assigned'] = 'LU2: Generative AI Applications in HR'
            # Day 2 content is likely related to LU4 if not otherwise assigned
            elif day_idx == 1:
                session['lu_assigned'] = 'LU4: Legal Consideration and the Future of Generative AI'
    
    # Rebuild the timetable LU by LU
    new_lesson_plan = []
    for day in lesson_plan:
        new_day = {'Day': day['Day'], 'Sessions': []}
        new_lesson_plan.append(new_day)
    
    # Copy all administrative sessions to the new plan first
    for day_idx, session_idx in admin_sessions:
        session = lesson_plan[day_idx]['Sessions'][session_idx]
        new_lesson_plan[day_idx]['Sessions'].append(deepcopy(session))
    
    # Sort administrative sessions within each day
    for day in new_lesson_plan:
        day['Sessions'].sort(key=lambda s: extract_time_for_sorting(s['Time']))
    
    # Process each LU in sequence, before final sessions
    current_day_idx = 0
    current_position = find_first_non_admin_position(new_lesson_plan[current_day_idx])
    
    # Find last position before final sessions on day 2
    final_day_idx = 1 if len(new_lesson_plan) > 1 else 0
    final_sessions_sorted = [(day_idx, session_idx) for day_idx, session_idx in final_sessions if day_idx == final_day_idx]
    final_sessions_sorted.sort(key=lambda x: extract_time_for_sorting(lesson_plan[x[0]]['Sessions'][x[1]]['Time']))
    
    # Find earliest final session position
    earliest_final_time = "2359"
    earliest_final_position = 999
    if final_sessions_sorted:
        day_idx, session_idx = final_sessions_sorted[0]
        session = lesson_plan[day_idx]['Sessions'][session_idx]
        earliest_final_time = extract_time_for_sorting(session['Time'])
        
        # Find where this would be in the new plan
        for i, s in enumerate(new_lesson_plan[day_idx]['Sessions']):
            if extract_time_for_sorting(s['Time']) >= earliest_final_time:
                earliest_final_position = i
                break
    
    # Process each LU in sequence
    for lu_idx, lu in enumerate(learning_units):
        lu_title = lu['LU_Title']
        lu_duration = lu_durations[lu_title]
        remaining_duration = lu_duration
        
        # Get topic and required duration for this LU
        lu_info = next((l for l in learning_units if l['LU_Title'] == lu_title), None)
        if not lu_info:
            continue
            
        topic_title = lu_info['Topics'][0]['Topic_Title'] if lu_info['Topics'] else f"Topic for {lu_title}"
        instructional_methods = ", ".join(lu_info['Instructional_Methods'])
        
        # Better balance - allocate approx 50% to topic and 50% to activities
        topic_duration = min(int(remaining_duration * 0.5), remaining_duration)
        remaining_duration -= topic_duration
        
        # Find where to insert this session
        day = new_lesson_plan[current_day_idx]
        
        # Format start time based on previous session end time or day start
        if current_position == 0 or len(day['Sessions']) == 0:
            start_time = "0930hrs"  # Day start time
        else:
            prev_session = day['Sessions'][current_position - 1]
            prev_end_time = re.search(r'- (\d{4})hrs', prev_session['Time']).group(1)
            start_time = prev_end_time
        
        start_hours, start_minutes = parse_time(start_time)
        
        # Calculate end time
        end_datetime = datetime(2023, 1, 1, start_hours, start_minutes) + timedelta(minutes=topic_duration)
        end_time = format_time(end_datetime.hour, end_datetime.minute)
        
        # Check if this would cross 1830 (END_TIME)
        if is_after_end_time(end_time):
            # Calculate how much time is available today
            end_of_day_hours, end_of_day_minutes = 18, 30  # 1830hrs
            available_minutes = (end_of_day_hours * 60 + end_of_day_minutes) - (start_hours * 60 + start_minutes)
            
            if available_minutes >= 30:  # At least 30 minutes left in the day
                # Use available time for first part
                first_part_duration = available_minutes
                remaining_duration += (topic_duration - first_part_duration)
                topic_duration = first_part_duration
                end_time = "1830hrs"
            else:
                # Move to next day
                current_day_idx += 1
                # If we run out of days, add a new day
                if current_day_idx >= len(new_lesson_plan):
                    new_day_idx = current_day_idx
                    new_day_number = int(new_lesson_plan[-1]['Day'].split(' ')[1]) + 1
                    new_day = {'Day': f'Day {new_day_number}', 'Sessions': []}
                    new_lesson_plan.append(new_day)
                
                current_position = find_first_non_admin_position(new_lesson_plan[current_day_idx])
                start_time = "0930hrs"  # Start at beginning of next day
                start_hours, start_minutes = 9, 30
                
                # Recalculate end time
                end_datetime = datetime(2023, 1, 1, start_hours, start_minutes) + timedelta(minutes=topic_duration)
                end_time = format_time(end_datetime.hour, end_datetime.minute)
                
                # Check if even on the new day it would cross end time
                if is_after_end_time(end_time):
                    # Further adjust to fit within the day
                    available_minutes = (18 * 60 + 30) - (9 * 60 + 30)  # 1830 - 0930 in minutes
                    first_part_duration = available_minutes
                    remaining_duration += (topic_duration - first_part_duration)
                    topic_duration = first_part_duration
                    end_time = "1830hrs"
        
        # Check if we would cross an administrative session
        next_admin_idx = find_next_admin_session(new_lesson_plan, current_day_idx, current_position)
        
        if next_admin_idx >= 0:
            next_admin_session = new_lesson_plan[current_day_idx]['Sessions'][next_admin_idx]
            next_admin_start = re.search(r'(\d{4})hrs', next_admin_session['Time']).group(1)
            next_hours, next_minutes = parse_time(next_admin_start)
            
            available_minutes = (next_hours * 60 + next_minutes) - (start_hours * 60 + start_minutes)
            
            if available_minutes < topic_duration:
                # Session would cross admin time, adjust duration or split
                if available_minutes >= 30:  # Minimum viable session length
                    # Use available time before admin session
                    adjusted_topic_duration = available_minutes
                    remaining_duration += (topic_duration - adjusted_topic_duration)
                    topic_duration = adjusted_topic_duration
                    
                    # Recalculate end time
                    end_datetime = datetime(2023, 1, 1, start_hours, start_minutes) + timedelta(minutes=topic_duration)
                    end_time = format_time(end_datetime.hour, end_datetime.minute)
                else:
                    # Skip to after admin session
                    next_admin_end = re.search(r'- (\d{4})hrs', next_admin_session['Time']).group(1)
                    start_time = next_admin_end
                    start_hours, start_minutes = parse_time(start_time)
                    current_position = next_admin_idx + 1
                    
                    # Recalculate end time
                    end_datetime = datetime(2023, 1, 1, start_hours, start_minutes) + timedelta(minutes=topic_duration)
                    end_time = format_time(end_datetime.hour, end_datetime.minute)
        
        # Create main topic session
        main_topic_session = {
            'Time': f"{start_time} - {end_time} ({format_duration(topic_duration)})",
            'instruction_title': topic_title,
            'bullet_points': lu_info['Topics'][0]['Bullet_Points'] if lu_info['Topics'] else [],
            'Instructional_Methods': instructional_methods,
            'Resources': "Slide page, TV, Whiteboard, Wi-Fi",
            'reference_line': 'Refer to some online references in Google Classroom LMS',
            'lu_assigned': lu_title,
            'is_topic': True,
            'is_activity': False
        }
        
        # Insert at current position
        new_lesson_plan[current_day_idx]['Sessions'].insert(current_position, main_topic_session)
        current_position += 1
        
        # Add activities to utilize the remaining duration
        if remaining_duration > 0:
            # Create a single activity session for better balance with the topic
            activity_duration = remaining_duration
            
            # Format start time based on previous session end time
            prev_session = new_lesson_plan[current_day_idx]['Sessions'][current_position - 1]
            prev_end_time = re.search(r'- (\d{4})hrs', prev_session['Time']).group(1)
            start_time = prev_end_time
            start_hours, start_minutes = parse_time(start_time)
            
            # Check if this would cross 1830 (END_TIME)
            end_datetime = datetime(2023, 1, 1, start_hours, start_minutes) + timedelta(minutes=activity_duration)
            end_time = format_time(end_datetime.hour, end_datetime.minute)
            
            if is_after_end_time(end_time):
                # Calculate how much time is available today
                end_of_day_hours, end_of_day_minutes = 18, 30  # 1830hrs
                available_minutes = (end_of_day_hours * 60 + end_of_day_minutes) - (start_hours * 60 + start_minutes)
                
                if available_minutes >= 30:  # At least 30 minutes left in the day
                    # Split activity into two parts
                    first_part_duration = available_minutes
                    
                    # Create first part of activity
                    activity_session = {
                        'Time': f"{start_time} - 1830hrs ({format_duration(first_part_duration)})",
                        'instruction_title': f"Activity: Case Study on {lu_title.split(':')[1].strip()} (Part 1)",
                        'bullet_points': [],
                        'Instructional_Methods': "Case Study, Group Discussion",
                        'Resources': "Worksheets, TV, Whiteboard, Wi-Fi",
                        'reference_line': 'Refer to some online references in Google Classroom LMS',
                        'lu_assigned': lu_title,
                        'is_topic': False,
                        'is_activity': True
                    }
                    
                    # Insert first activity
                    new_lesson_plan[current_day_idx]['Sessions'].insert(current_position, activity_session)
                    current_position += 1
                    
                    # Move to next day for second part
                    current_day_idx += 1
                    # If we run out of days, add a new day
                    if current_day_idx >= len(new_lesson_plan):
                        new_day_idx = current_day_idx
                        new_day_number = int(new_lesson_plan[-1]['Day'].split(' ')[1]) + 1
                        new_day = {'Day': f'Day {new_day_number}', 'Sessions': []}
                        new_lesson_plan.append(new_day)
                    
                    current_position = find_first_non_admin_position(new_lesson_plan[current_day_idx])
                    start_time = "0930hrs"  # Start at beginning of next day
                    
                    # Create second part of activity with remaining duration
                    remaining_activity_duration = activity_duration - first_part_duration
                    end_datetime = datetime(2023, 1, 1, 9, 30) + timedelta(minutes=remaining_activity_duration)
                    end_time = format_time(end_datetime.hour, end_datetime.minute)
                    
                    # Check if even on the new day it would cross end time
                    if is_after_end_time(end_time):
                        # Further split if needed
                        available_minutes = (18 * 60 + 30) - (9 * 60 + 30)  # 1830 - 0930 in minutes
                        if remaining_activity_duration > available_minutes:
                            remaining_activity_duration = available_minutes
                            end_time = "1830hrs"
                    
                    activity_session = {
                        'Time': f"{start_time} - {end_time} ({format_duration(remaining_activity_duration)})",
                        'instruction_title': f"Activity: Case Study on {lu_title.split(':')[1].strip()} (Part 2)",
                        'bullet_points': [],
                        'Instructional_Methods': "Case Study, Group Discussion, Practice",
                        'Resources': "Worksheets, TV, Whiteboard, Wi-Fi",
                        'reference_line': 'Refer to some online references in Google Classroom LMS',
                        'lu_assigned': lu_title,
                        'is_topic': False,
                        'is_activity': True
                    }
                    
                    # Insert second activity
                    new_lesson_plan[current_day_idx]['Sessions'].insert(current_position, activity_session)
                    current_position += 1
                else:
                    # Move entire activity to next day
                    current_day_idx += 1
                    # If we run out of days, add a new day
                    if current_day_idx >= len(new_lesson_plan):
                        new_day_idx = current_day_idx
                        new_day_number = int(new_lesson_plan[-1]['Day'].split(' ')[1]) + 1
                        new_day = {'Day': f'Day {new_day_number}', 'Sessions': []}
                        new_lesson_plan.append(new_day)
                    
                    current_position = find_first_non_admin_position(new_lesson_plan[current_day_idx])
                    start_time = "0930hrs"  # Start at beginning of next day
                    
                    # Recalculate end time
                    end_datetime = datetime(2023, 1, 1, 9, 30) + timedelta(minutes=activity_duration)
                    end_time = format_time(end_datetime.hour, end_datetime.minute)
                    
                    # Check if even on the new day it would cross end time
                    if is_after_end_time(end_time):
                        # Further adjust to fit within the day
                        available_minutes = (18 * 60 + 30) - (9 * 60 + 30)  # 1830 - 0930 in minutes
                        activity_duration = available_minutes
                        end_time = "1830hrs"
                    
                    activity_session = {
                        'Time': f"{start_time} - {end_time} ({format_duration(activity_duration)})",
                        'instruction_title': f"Activity: Case Study and Practice on {lu_title.split(':')[1].strip()}",
                        'bullet_points': [],
                        'Instructional_Methods': "Case Study, Group Discussion, Practice",
                        'Resources': "Worksheets, TV, Whiteboard, Wi-Fi",
                        'reference_line': 'Refer to some online references in Google Classroom LMS',
                        'lu_assigned': lu_title,
                        'is_topic': False,
                        'is_activity': True
                    }
                    
                    # Insert activity
                    new_lesson_plan[current_day_idx]['Sessions'].insert(current_position, activity_session)
                    current_position += 1
            else:
                # Check if we would cross an administrative session
                next_admin_idx = find_next_admin_session(new_lesson_plan, current_day_idx, current_position)
                
                if next_admin_idx >= 0:
                    next_admin_session = new_lesson_plan[current_day_idx]['Sessions'][next_admin_idx]
                    next_admin_start = re.search(r'(\d{4})hrs', next_admin_session['Time']).group(1)
                    next_hours, next_minutes = parse_time(next_admin_start)
                    
                    available_minutes = (next_hours * 60 + next_minutes) - (start_hours * 60 + start_minutes)
                    
                    if available_minutes < activity_duration:
                        # Split activity if needed
                        if available_minutes >= 30:
                            # Create first part of activity
                            first_activity_duration = available_minutes
                            
                            # Calculate end time for first part
                            end_datetime = datetime(2023, 1, 1, start_hours, start_minutes) + timedelta(minutes=first_activity_duration)
                            end_time = format_time(end_datetime.hour, end_datetime.minute)
                            
                            # Create first activity session
                            activity_session = {
                                'Time': f"{start_time} - {end_time} ({format_duration(first_activity_duration)})",
                                'instruction_title': f"Activity: Discussion on {lu_title.split(':')[1].strip()}",
                                'bullet_points': [],
                                'Instructional_Methods': "Group Discussion, Peer Sharing",
                                'Resources': "Worksheets, TV, Whiteboard, Wi-Fi",
                                'reference_line': 'Refer to some online references in Google Classroom LMS',
                                'lu_assigned': lu_title,
                                'is_topic': False,
                                'is_activity': True
                            }
                            
                            # Insert first activity
                            new_lesson_plan[current_day_idx]['Sessions'].insert(current_position, activity_session)
                            current_position += 1
                            
                            # Prepare for second part after admin session
                            activity_duration -= first_activity_duration
                            next_admin_end = re.search(r'- (\d{4})hrs', next_admin_session['Time']).group(1)
                            start_time = next_admin_end
                            start_hours, start_minutes = parse_time(start_time)
                            current_position = next_admin_idx + 1
                        else:
                            # Skip to after admin session
                            next_admin_end = re.search(r'- (\d{4})hrs', next_admin_session['Time']).group(1)
                            start_time = next_admin_end
                            start_hours, start_minutes = parse_time(start_time)
                            current_position = next_admin_idx + 1
                
                # Calculate end time
                end_datetime = datetime(2023, 1, 1, start_hours, start_minutes) + timedelta(minutes=activity_duration)
                end_time = format_time(end_datetime.hour, end_datetime.minute)
                
                # Create activity session
                activity_session = {
                    'Time': f"{start_time} - {end_time} ({format_duration(activity_duration)})",
                    'instruction_title': f"Activity: Case Study and Practice on {lu_title.split(':')[1].strip()}",
                    'bullet_points': [],
                    'Instructional_Methods': "Case Study, Group Discussion, Practice",
                    'Resources': "Worksheets, TV, Whiteboard, Wi-Fi",
                    'reference_line': 'Refer to some online references in Google Classroom LMS',
                    'lu_assigned': lu_title,
                    'is_topic': False,
                    'is_activity': True
                }
                
                # Insert activity
                new_lesson_plan[current_day_idx]['Sessions'].insert(current_position, activity_session)
                current_position += 1
    
        # Calculate total duration of all final sessions
    final_sessions_duration = 0
    final_sessions_data = []
    
    for day_idx, session_idx in final_sessions:
        session = lesson_plan[day_idx]['Sessions'][session_idx]
        duration = extract_duration_minutes(session['Time'])
        final_sessions_duration += duration
        final_sessions_data.append({
            'session': deepcopy(session),
            'duration': duration
        })
    
    # Sort final sessions - TRAQOM survey before assessments if needed
    final_sessions_data.sort(key=lambda x: 0 if 'survey' in x['session']['instruction_title'].lower() or 'feedback' in x['session']['instruction_title'].lower() else 1)
    
    # Find the last day in the timetable
    last_day_idx = len(new_lesson_plan) - 1
    
    # Calculate start time by working backward from 1830hrs
    end_time = "1830hrs"
    end_hours, end_minutes = 18, 30
    
    # Calculate when final sessions should start
    start_datetime = datetime(2023, 1, 1, end_hours, end_minutes) - timedelta(minutes=final_sessions_duration)
    start_hours, start_minutes = start_datetime.hour, start_datetime.minute
    final_sessions_start_time = format_time(start_hours, start_minutes)
    
    # Add short end-of-day recaps (5 min) for each day EXCEPT the final assessment day
    for day_idx, day in enumerate(new_lesson_plan):
        # Skip the day with final assessments
        if day_idx == last_day_idx and final_sessions_data:
            continue
            
        # Check if there's already a session ending at 1830
        has_session_ending_at_1830 = False
        for session in day['Sessions']:
            end_time_match = re.search(r'- (\d{4})hrs', session['Time'])
            if end_time_match and end_time_match.group(1) == "1830":
                has_session_ending_at_1830 = True
                break
                
        if not has_session_ending_at_1830:
            # Create a 5-minute recap session at the end of the day
            recap_session = {
                'Time': f"1825hrs - 1830hrs (5 mins)",
                'instruction_title': 'Recap All Contents and Close',
                'bullet_points': ['Summary of key learning points', 'Q&A'],
                'Instructional_Methods': 'Lecture, Group Discussion',
                'Resources': 'Slide page 16, TV, Whiteboard, Wi-Fi',
                'reference_line': 'Refer to course slides and notes'
            }
            day['Sessions'].append(recap_session)
    
    # Make sure all regular content ends before final sessions start
    # Go through each day and adjust sessions that would overlap with finals
    for day_idx, day in enumerate(new_lesson_plan):
        if day_idx == last_day_idx:
            # Find any sessions that would extend past the final sessions start time
            sessions_to_remove = []
            for i, session in enumerate(day['Sessions']):
                if session.get('is_admin', False) or session.get('is_final', True):
                    continue
                    
                session_end_match = re.search(r'- (\d{4})hrs', session['Time'])
                if session_end_match:
                    session_end_time = session_end_match.group(1)
                    if extract_time_for_sorting(session_end_time) > extract_time_for_sorting(final_sessions_start_time):
                        # This session would overlap with finals - truncate it
                        session_start_match = re.search(r'(\d{4})hrs', session['Time'])
                        if session_start_match:
                            session_start_time = session_start_match.group(1)
                            session_start_hours, session_start_minutes = parse_time(session_start_time)
                            
                            # Calculate new duration
                            available_minutes = (start_hours * 60 + start_minutes) - (session_start_hours * 60 + start_minutes)
                            
                            if available_minutes >= 30:  # At least 30 minutes for a meaningful session
                                # Update the session to end at the start of final sessions
                                session['Time'] = f"{session_start_time} - {final_sessions_start_time} ({format_duration(available_minutes)})"
                            else:
                                # Session is too short after truncation - mark for removal
                                sessions_to_remove.append(i)
            
            # Remove sessions that are too short (in reverse order to avoid index shifts)
            for i in sorted(sessions_to_remove, reverse=True):
                day['Sessions'].pop(i)
    
    # On assessment day, add a longer recap session if there's extra time between last LU and final sessions
    if final_sessions_data:
        assessment_day = new_lesson_plan[last_day_idx]
        
        # Find the last non-administrative session before final sessions start
        last_content_idx = -1
        last_content_end_time = None
        
        for i, session in enumerate(assessment_day['Sessions']):
            # Skip administrative sessions
            if session.get('is_admin', False):
                continue
                
            # Skip final sessions
            if session.get('is_final', False):
                continue
                
            # Check if this session ends before final sessions start
            session_end_match = re.search(r'- (\d{4})hrs', session['Time'])
            if session_end_match:
                session_end_time = session_end_match.group(1)
                if extract_time_for_sorting(session_end_time) <= extract_time_for_sorting(final_sessions_start_time):
                    last_content_idx = i
                    last_content_end_time = session_end_time
        
        # If we found the last content session and there's a gap before final sessions
        if last_content_idx >= 0 and last_content_end_time:
            # Calculate the gap duration in minutes
            last_end_hours, last_end_minutes = parse_time(last_content_end_time)
            final_start_hours, final_start_minutes = parse_time(final_sessions_start_time)
            
            gap_minutes = ((final_start_hours * 60 + final_start_minutes) - 
                          (last_end_hours * 60 + last_end_minutes))
            
            # If there's a significant gap (at least 15 minutes), add a recap session
            if gap_minutes >= 15:
                comprehensive_recap_session = {
                    'Time': f"{last_content_end_time} - {final_sessions_start_time} ({format_duration(gap_minutes)})",
                    'instruction_title': 'Comprehensive Course Review and Assessment Preparation',
                    'bullet_points': [
                        "Review of key concepts from all Learning Units",
                        "Integration of learning across different topics",
                        "Application scenarios and practical insights",
                        "Assessment preparation and clarification of doubts",
                        "Final Q&A session"
                    ],
                    'Instructional_Methods': "Lecture, Group Discussion, Q&A",
                    'Resources': "Slide pages, TV, Whiteboard, Wi-Fi",
                    'reference_line': 'Refer to all course materials'
                }
                
                # Insert right after the last content session
                assessment_day['Sessions'].insert(last_content_idx + 1, comprehensive_recap_session)
    
    # Remove any existing final sessions since we'll add them back in the correct order
    for day in new_lesson_plan:
        day['Sessions'] = [s for s in day['Sessions'] if not s.get('is_final', False)]
    
    # First sort all sessions in each day by time (for all non-final sessions)
    for day in new_lesson_plan:
        day['Sessions'].sort(key=lambda s: extract_time_for_sorting(s['Time']))
    
    # Add final sessions in order to the last day - these must always be last
    current_time = final_sessions_start_time
    final_sessions_to_add = []
    
    for session_data in final_sessions_data:
        session = session_data['session']
        duration = session_data['duration']
        
        # Calculate end time for this session
        current_hours, current_minutes = parse_time(current_time)
        end_datetime = datetime(2023, 1, 1, current_hours, current_minutes) + timedelta(minutes=duration)
        session_end_time = format_time(end_datetime.hour, end_datetime.minute)
        
        # Update session time
        session['Time'] = f"{current_time} - {session_end_time} ({format_duration(duration)})"
        session['is_final'] = True  # Mark as final session for sorting
        
        # Add to collection (don't add directly to day yet)
        final_sessions_to_add.append(session)
        
        # Next session starts where this one ends
        current_time = session_end_time
    
    # Now add all final sessions to the end of the last day
    new_lesson_plan[last_day_idx]['Sessions'].extend(final_sessions_to_add)
    
    # Final cleanup
    for day in new_lesson_plan:
        for session in day['Sessions']:
            if 'lu_assigned' in session:
                del session['lu_assigned']
            if 'is_admin' in session:
                del session['is_admin']
            if 'is_final' in session:
                del session['is_final']
            if 'is_topic' in session:
                del session['is_topic']
            if 'is_activity' in session:
                del session['is_activity']
    # Update the lesson plan in the context
    context['lesson_plan'] = new_lesson_plan
    return context

# Helper functions for the timetable optimization
def find_first_non_admin_position(day):
    """Find the position after the first attendance session"""
    for i, session in enumerate(day['Sessions']):
        if 'attendance' in session['instruction_title'].lower():
            return i + 1
    return 0

def find_next_admin_session(lesson_plan, day_idx, current_position):
    """Find the next administrative session (break, lunch, etc.)"""
    day = lesson_plan[day_idx]
    for i in range(current_position, len(day['Sessions'])):
        session = day['Sessions'][i]
        if session.get('is_admin', False):
            return i
    return -1


def parse_cp_document(uploaded_file):
    """
    Parses a Course Proposal (CP) document (UploadedFile) and returns its content as Markdown text,
    trimmed based on the document type using regex patterns.

    For Word CP (.docx):
      - Excludes everything before a line matching "Part 1" and "Particulars of Course"
      - Excludes everything after a line matching "Part 4" and "Facilities and Resources"
    
    For Excel CP (.xlsx):
      - Excludes everything before a line matching "1 - Course Particulars"
      - Excludes everything after a line matching "3 - Summary"

    Args:
        uploaded_file (UploadedFile): The file uploaded via st.file_uploader.

    Returns:
        str: A trimmed Markdown string containing the parsed document content.
    """
    # Write the uploaded file to a temporary file.
    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp:
        tmp.write(uploaded_file.read())
        temp_file_path = tmp.name

    try:
        # Set up parser for markdown result
        parser = LlamaParse(result_type="markdown", api_key=st.secrets["LLAMA_CLOUD_API_KEY"], 
        system_prompt_append=(
        "**IMPORTANT FOLLOW THE RULES OR YOU WILL FAIL**"
        "1.Extract the entire content of the document without omitting any text, including line breaks, paragraphs, and list items, in the exact order they appear. "
        "2.Parse tables into Markdown tables with the following specific rules:"
        "3. Ensure that the instructional methods column under curriculum table is flattened, meaning that if it contains multiple lines, they should be concatenated into a single line. "
        "4. When a table cell contains multiple lines, especially in the 'Instructional Methods' column, concatenate them as a single string, separated by commas. "
        "5. Never preserve newlines inside table cells in the output Markdown. "
        "6. For the curriculum table with columns like 'S/N', 'LUs', 'LOs*', and 'Instructional Methods', pay special attention to flatten all content in each cell. "
        "7. Ensure the output is a faithful and complete Markdown representation of the document."
        ))
        
    
        # Determine the file extension for mapping
        ext = os.path.splitext(temp_file_path)[1].lower()
        file_extractor = {ext: parser}
    
        # Use SimpleDirectoryReader to load and parse the file
        documents = SimpleDirectoryReader(input_files=[temp_file_path], file_extractor=file_extractor).load_data()
        print("Lllama Parser")
        print(documents)
    
        # Concatenate the parsed text from each Document object into a single Markdown string
        markdown_text = "\n\n".join(doc.text for doc in documents)
    
        # Set up regex patterns based on file extension
        if ext == ".docx":
            start_pattern = re.compile(r"Part\s*1.*?Particulars\s+of\s+Course", re.IGNORECASE)
            end_pattern = re.compile(r"Part\s*4.*?Facilities\s+and\s+Resources", re.IGNORECASE)
        elif ext == ".xlsx":
            start_pattern = re.compile(r"1\s*-\s*Course\s*Particulars", re.IGNORECASE)
            end_pattern = re.compile(r"4\s*-\s*Declarations", re.IGNORECASE)
        else:
            start_pattern = None
            end_pattern = None
    
        # If both patterns exist, search for the matches and trim the text
        if start_pattern and end_pattern:
            start_match = start_pattern.search(markdown_text)
            end_match = end_pattern.search(markdown_text)
            if start_match and end_match and end_match.start() > start_match.start():
                markdown_text = markdown_text[start_match.start():end_match.start()].strip()
    
    finally:
        # Clean up the temporary file
        os.remove(temp_file_path)
    
    return markdown_text

############################################################
# 2. Interpret Course Proposal Data
############################################################
async def interpret_cp(raw_data: dict, model_client: OpenAIChatCompletionClient) -> dict:
    """
    Interprets and extracts structured data from a raw Course Proposal (CP) document.

    This function processes raw CP data using an AI model to extract 
    structured information such as course details, learning units, topics, 
    assessment methods, and instructional methods.

    Args:
        raw_data (dict): 
            The unstructured data extracted from the CP document.
        model_client (OpenAIChatCompletionClient): 
            The AI model client used for structured data extraction.

    Returns:
        dict: 
            A structured dictionary containing course details.

    Raises:
        Exception: 
            If the AI-generated response does not contain the expected fields.
    """

    # Interpreter Agent with structured output enforcement
    interpreter = AssistantAgent(
        name="Interpreter",
        model_client=model_client,
        system_message=f"""
        You are an AI assistant that helps extract specific information from a JSON object containing a Course Proposal Form (CP). Your task is to interpret the JSON data, regardless of its structure, and extract the required information accurately.

        ---
        
        **Task:** Extract the following information from the provided JSON data:

        ### Part 1: Particulars of Course

        - Name of Organisation
        - Course Title
        - TSC Title
        - TSC Code
        - Total Training Hours/ Total Instructional Duration (calculated as the sum of Classroom Facilitation, Workplace Learning: On-the-Job (OJT), Practicum, Practical, E-learning: Synchronous and Asynchronous), formatted with units (e.g., "30 hrs", "1 hr")
        - Total Assessment Hours/ Total Assessment Duration, formatted with units (e.g., "2 hrs")
        - Total Course Duration Hours, formatted with units (e.g., "42 hrs")

        ### Part 3: Curriculum Design

        From the Learning Units and Topics Table:

        For each Learning Unit (LU):
        - Learning Unit Title (include the "LUx: " prefix)
        - **Learning Unit Duration (in hours or minutes, as specified in the CP)**
        - Topics Covered Under Each LU:
        - For each Topic:
            - **Topic_Title** (include the "Topic x: " prefix and the associated K and A statements in parentheses)
            - **Bullet_Points** (a list of bullet points under the topic; remove any leading bullet symbols such as "-" so that only the content remains)
        - Learning Outcomes (LOs) (include the "LOx: " prefix for each LO)
        - Numbering and Description for the "K" (Knowledge) Statements (as a list of dictionaries with keys "K_number" and "Description")
        - Numbering and Description for the "A" (Ability) Statements (as a list of dictionaries with keys "A_number" and "Description")
        - **Assessment_Methods** (a list of assessment method abbreviations; e.g., ["WA-SAQ", "CS"]). Note: If the CP contains the term "Written Exam", output it as "Written Assessment - Short Answer Questions". If it contains "Practical Exam", output it as "Practical Performance".
        - **Duration Calculation:** When extracting the duration for each assessment method:
            1. If the extracted duration is not exactly 0.5 or a whole number (e.g., 0.5, 1, 2, etc.), interpret it as minutes.
            2. If duplicate entries for the same assessment method occur within the same LU, sum their durations to obtain a total duration.
            3. For CPs in Excel format, under 3 - Summary sheet, the duration appears in the format "(Assessor-to-Candidate Ratio, duration)"—for example, "Written Exam (1:20, 20)" means 20 minutes, and "Others: Case Study (1:20, 25)" appearing twice should result in a total of 50 minutes for Case Study.       
        - **Instructional_Methods** (a list of instructional method abbreviations or names)

        ### Part E: Details of Assessment Methods Proposed

        For each Assessment Method in the CP, extract:
        - **Assessment_Method** (always use the full term, e.g., "Written Assessment - Short Answer Questions", "Practical Performance", "Case Study", "Oral Questioning", "Role Play")
        - **Method_Abbreviation** (if provided in parentheses or generated according to the rules)
        - **Total_Delivery_Hours** (formatted with units, e.g., "1 hr")
        - **Assessor_to_Candidate_Ratio** (a list of minimum and maximum ratios, e.g., ["1:3 (Min)", "1:5 (Max)"])
        
        **Additionally, if the CP explicitly provides the following fields, extract them. Otherwise, do not include them in the final output:**
        - **Type_of_Evidence**  
        - For PP and CS assessment methods, the evidence may be provided as a dictionary where keys are LO identifiers (e.g., "LO1", "LO2", "LO3") and values are the corresponding evidence text. In that case, convert the dictionary into a list of dictionaries with keys `"LO"` and `"Evidence"`.  
        - If the evidence is already provided as a list (for example, a list of strings or a list of dictionaries), keep it as is.
        - **Manner_of_Submission** (as a list, e.g., ["Submission 1", "Submission 2"])
        - **Marking_Process** (as a list, e.g., ["Process 1", "Process 2"])
        - **Retention_Period**: **Extract the complete retention description exactly as provided in the CP.**
        - **No_of_Role_Play_Scripts** (only if the assessment method is Role Play and this information is provided)

        ---
        
        **Instructions:**
        
        - Carefully parse the JSON data and locate the sections corresponding to each part.
        - Even if the JSON structure changes, use your understanding to find and extract the required information.
        - Ensure that the `Topic_Title` includes the "Topic x: " prefix and the associated K and A statements in parentheses exactly as they appear.
        - For Learning Outcomes (LOs), always include the "LOx: " prefix (where x is the number).
        - Present the extracted information in a structured JSON format where keys correspond exactly to the placeholders required for the Word document template.
        - Ensure all extracted information is normalized by:
            - Replacing en dashes (–) and em dashes (—) with hyphens (-)
            - Converting curly quotes (“ ”) to straight quotes (")
            - Replacing other non-ASCII characters with their closest ASCII equivalents.
        - **Time fields** must include units (e.g., "40 hrs", "1 hr", "2 hrs").
        - For `Assessment_Methods`, always use the abbreviations (e.g., WA-SAQ, PP, CS, OQ, RP) as per the following rules:
            1. Use the abbreviation provided in parentheses if available.
            2. Otherwise, generate an abbreviation by taking the first letters of the main words (ignoring articles/prepositions) and join with hyphens.
            3. For methods containing "Written Assessment", always prefix with "WA-".
            4. If duplicate or multiple variations exist, use the standard abbreviation.
        - **Important:** Verify that the sum of `Total_Delivery_Hours` for all assessment methods equals the `Total_Assessment_Hours`. If individual delivery hours for assessment methods are not specified, divide the `Total_Assessment_Hours` equally among them.
        - For bullet points in each topic, ensure that the number of bullet points exactly matches those in the CP. Re-extract if discrepancies occur.
        - **If the same K or A statement (same numbering and description) appears multiple times within the same LU, keep only one instance. If the same K or A statement appears in different LUs, keep it as it is.**
        - Do not include any extraneous information or duplicate entries.

        Generate structured output matching this schema:
        {json.dumps(CourseData.model_json_schema(), indent=2)}
        """,
    )

    agent_task = f"""
    Please extract and structure the following data: {raw_data}.
    **Return the extracted information as a complete JSON dictionary containing the specified fields. Do not truncate or omit any data. Include all fields and their full content. Do not use '...' or any placeholders to replace data.**
    Simply return the JSON dictionary object directly.
    """

    # Process sample input
    response = await interpreter.on_messages(
        [TextMessage(content=agent_task, source="user")], CancellationToken()
    )
    if not response or not response.chat_message:
        return "No content found in the agent's last message."
    # print(response.chat_message.content)
    # return response.chat_message.content

    context = parse_json_content(response.chat_message.content)
    return context

# Streamlit App
def app():
    """
    Streamlit web application for generating courseware documents.

    This function serves as the entry point for the user interface,
    allowing users to upload a Course Proposal document, select 
    their organization, and generate various courseware documents.

    The app guides users through:
    - Uploading a Course Proposal (CP) document.
    - Selecting an organization from a predefined list.
    - Uploading an optional updated Skills Framework (SFw) dataset.
    - Selecting documents to generate (Learning Guide, Lesson Plan, etc.).
    - Processing and downloading the generated documents.

    Raises:
        ValueError: 
            If required input fields are missing.
        Exception: 
            If any step in the document generation process fails.
    """

    st.title("📄 Courseware Document Generator")
    
    # ================================================================
    # MODEL SELECTION FEATURE
    # ================================================================
    st.subheader("Model Selection")
    model_choice = st.selectbox(
        "Select LLM Model:",
        options=list(MODEL_CHOICES.keys()),
        index=0 # Select Default
    )
    st.session_state['selected_model'] = model_choice

    # ================================================================
    # Step 1: Upload Course Proposal (CP) Document
    # ================================================================
    st.subheader("Step 1: Upload Course Proposal (CP) Document")
    cp_file = st.file_uploader("Upload Course Proposal (CP) Document", type=["docx, xlsx"])

    # ================================================================
    # Step 2: Select Name of Organisation
    # ================================================================
    # Create a modal instance with a unique key and title
    crud_modal = Modal(key="crud_modal", title="Manage Organisations")

    st.subheader("Step 2: Enter Relevant Details")
    tgs_course_code = st.text_input("Enter TGS Course Code", key="tgs_course_code", placeholder="e.g., TGS-2023039181")

    col1, col2 = st.columns([0.8, 0.2], vertical_alignment="center")
    # Load organisations from JSON using the utility function
    org_list = load_organizations()
    org_names = [org["name"] for org in org_list] if org_list else []
    with col1:
        if org_names:
            selected_org = st.selectbox("Select Name of Organisation", org_names, key="org_dropdown_main")
        else:
            selected_org = st.selectbox("Select Name of Organisation", [], key="org_dropdown_main")
            st.warning("No organisations found. Click 'Manage' to add organisations.")

    with col2:
        # Wrap the Manage button in a div that uses flexbox for vertical centering.
        st.markdown("<br/>", unsafe_allow_html=True)
        if st.button("Manage", key="manage_button", use_container_width=True):
            crud_modal.open()

    # ---------------------------
    # Modal: CRUD Interface
    # ---------------------------
    if crud_modal.is_open():
        with crud_modal.container():
            
            # ---- Add New Organisation Form ----
            st.write("#### Add New Organisation")
            with st.form("new_org_form"):
                new_name = st.text_input("Organisation Name", key="new_org_name")
                new_uen = st.text_input("UEN", key="new_org_uen")
                # Use file uploader for the logo instead of a text input
                new_logo_file = st.file_uploader("Upload Logo (optional)", type=["png", "jpg", "jpeg"], key="new_org_logo_file")
                new_submitted = st.form_submit_button("Add Organisation")
                if new_submitted:
                    logo_path = None
                    if new_logo_file is not None:
                        # Construct a safe filename based on the organisation name and file extension
                        _, ext = os.path.splitext(new_logo_file.name)
                        safe_filename = new_name.lower().replace(" ", "_") + ext
                        save_path = os.path.join("Courseware", "utils", "logo", safe_filename)
                        with open(save_path, "wb") as f:
                            f.write(new_logo_file.getvalue())
                        logo_path = save_path
                    new_org = Organization(name=new_name, uen=new_uen, logo=logo_path)
                    add_organization(new_org)
                    st.success(f"Organisation '{new_name}' added.")
                    st.rerun()
            
            # ---- Display Existing Organisations with Edit/Delete Buttons ----
            st.write("#### Existing Organisations")
            org_list = load_organizations()  # Refresh the list

            # Table header
            col_sno, col_name, col_uen, col_logo, col_edit, col_delete = st.columns([1, 3, 2, 2, 1, 2])
            col_sno.markdown("**SNo**")
            col_name.markdown("**Name**")
            col_uen.markdown("**UEN**")
            col_logo.markdown("**Logo**")
            col_edit.markdown("**Edit**")
            col_delete.markdown("**Delete**")

            # Table rows
            for display_idx, org in enumerate(org_list, start=1):
                # The actual index in the list is display_idx - 1
                real_index = display_idx - 1

                row_sno, row_name, row_uen, row_logo, row_edit, row_delete = st.columns([1, 3, 2, 2, 1, 2])
                row_sno.write(display_idx)
                row_name.write(org["name"])
                row_uen.write(org["uen"])
                
                if org["logo"] and os.path.exists(org["logo"]):
                    row_logo.image(org["logo"], width=70)
                else:
                    row_logo.write("No Logo")

                # Edit/Delete Buttons
                if row_edit.button("Edit", key=f"edit_{display_idx}", type="secondary"):
                    st.session_state["org_edit_index"] = real_index
                    st.rerun()
                if row_delete.button("Delete", key=f"delete_{display_idx}", type="primary"):
                    if org["logo"] and os.path.exists(org["logo"]):
                        os.remove(org["logo"])
                    delete_organization(real_index)
                    st.success(f"Organisation '{org['name']}' deleted.")
                    st.rerun()

            # ---- Edit Organisation Form (if a row is selected for editing) ----
            if "org_edit_index" in st.session_state:
                edit_index = st.session_state["org_edit_index"]
                org_to_edit = load_organizations()[edit_index]
                st.write(f"#### Edit Organisation: {org_to_edit['name']}")
                with st.form("edit_org_form"):
                    edited_name = st.text_input("Organisation Name", value=org_to_edit["name"], key="edited_name")
                    edited_uen = st.text_input("UEN", value=org_to_edit["uen"], key="edited_uen")
                    # File uploader for updating the logo image
                    edited_logo_file = st.file_uploader("Upload Logo (optional)", type=["png", "jpg", "jpeg"], key="edited_logo_file")
                    edit_submitted = st.form_submit_button("Update Organisation")
                    if edit_submitted:
                        logo_path = org_to_edit.get("logo", None)
                        if edited_logo_file is not None:
                            _, ext = os.path.splitext(edited_logo_file.name)
                            safe_filename = edited_name.lower().replace(" ", "_") + ext
                            save_path = os.path.join("Courseware", "utils", "logo", safe_filename)
                            with open(save_path, "wb") as f:
                                f.write(edited_logo_file.getvalue())
                            logo_path = save_path
                        updated_org = Organization(name=edited_name, uen=edited_uen, logo=logo_path)
                        update_organization(edit_index, updated_org)
                        st.success(f"Organisation '{edited_name}' updated.")
                        del st.session_state["org_edit_index"]
                        st.rerun()

    # ================================================================
    # Step 3 (Optional): Upload Updated SFW Dataset
    # ================================================================
    st.subheader("Step 3 (Optional): Upload Updated Skills Framework (SFw) Dataset")
    sfw_file = st.file_uploader("Upload Updated SFw Dataset (Excel File)", type=["xlsx"])
    if sfw_file:
        sfw_data_dir = save_uploaded_file(sfw_file, 'input/dataset')
        st.success(f"Updated SFw dataset saved to {sfw_data_dir}")
    else:
        sfw_data_dir = "Courseware/input/dataset/Sfw_dataset-2022-03-30 copy.xlsx"

    # ================================================================
    # Step 4: Select Document(s) to Generate using Checkboxes
    # ================================================================
    st.subheader("Step 4: Select Document(s) to Generate")
    generate_lg = st.checkbox("Learning Guide (LG)")
    generate_ap = st.checkbox("Assessment Plan (AP)")
    generate_lp = st.checkbox("Lesson Plan (LP)")
    generate_fg = st.checkbox("Facilitator's Guide (FG)")

    # ================================================================
    # Step 5: Generate Documents
    # ================================================================
    if st.button("Generate Documents"):
        if cp_file is not None and selected_org:
            # Reset previous output document paths
            st.session_state['lg_output'] = None
            st.session_state['ap_output'] = None
            st.session_state['asr_output'] = None
            st.session_state['lp_output'] = None
            st.session_state['fg_output'] = None
            # Use the selected model configuration for all autogen agents
            selected_config = get_model_config(st.session_state['selected_model'])
            api_key = selected_config["config"].get("api_key")
            if not api_key:
                st.error("API key for the selected model is not provided.")
                return
            model_name = selected_config["config"]["model"]
            temperature = selected_config["config"].get("temperature", 0)
            base_url = selected_config["config"].get("base_url", None)

            # Extract model_info from the selected configuration (if provided)
            model_info = selected_config["config"].get("model_info", None)

            # Conditionally set response_format: use structured output only for valid OpenAI models.
            if st.session_state['selected_model'] in ["DeepSeek-V3", "Gemini-Pro-2.5-Exp-03-25"]:
                cp_response_format = None  # DeepSeek and Gemini might not support structured output this way.
                lp_response_format = None
            else:
                cp_response_format = CourseData  # For structured CP extraction
                lp_response_format = LessonPlan  # For timetable generation

            openai_struct_model_client = OpenAIChatCompletionClient(
                model=model_name,
                api_key=api_key,
                temperature=temperature,
                base_url=base_url,
                response_format=cp_response_format,  # Only set for valid OpenAI models
                model_info=model_info,
            )

            timetable_openai_struct_model_client = OpenAIChatCompletionClient(
                model=model_name,
                api_key=api_key,
                temperature=temperature,
                base_url=base_url,
                response_format=lp_response_format,
                model_info=model_info,
            )

            openai_model_client = OpenAIChatCompletionClient(
                model=model_name,
                api_key=api_key,
                temperature=temperature,
                base_url=base_url,
                model_info=model_info,
            )

            # Step 1: Parse the CP document
            try:
                with st.spinner('Parsing the Course Proposal...'):
                    raw_data = parse_cp_document(cp_file)
                    #New!Check output
                    print("Raw data:")
                    print(raw_data)
            except Exception as e:
                st.error(f"Error parsing the Course Proposal: {e}")
                return
            
            try:
                with st.spinner('Extracting Information from Course Proposal...'):
                    context = asyncio.run(interpret_cp(raw_data=raw_data, model_client=openai_struct_model_client))
                    print("Extracted Context:")
                    print(context)

            except Exception as e:
                st.error(f"Error extracting Course Proposal: {e}")
                return

            # After obtaining the context
            if context:
                # Step 2: Add the current date to the raw_data
                current_datetime = datetime.now()
                current_date = current_datetime.strftime("%d %b %Y")
                year = current_datetime.year
                context["Date"] = current_date
                context["Year"] = year
                # Find the selected organisation UEN in the organisation's record
                selected_org_data = next((org for org in org_list if org["name"] == selected_org), None)
                if selected_org_data:
                    context["UEN"] = selected_org_data["uen"]

                tgs_course_code = st.session_state.get("tgs_course_code", "")
                context["TGS_Ref_No"] = tgs_course_code

                st.session_state['context'] = context  # Store context in session state

                # Generate Learning Guide
                if generate_lg:
                    try:
                        with st.spinner('Generating Learning Guide...'):
                            lg_output = generate_learning_guide(context, selected_org, openai_model_client)
                        if lg_output:
                            st.success(f"Learning Guide generated: {lg_output}")
                            st.session_state['lg_output'] = lg_output  # Store output path in session state
                    except Exception as e:
                        st.error(f"Error generating Learning Guide: {e}")

                # Generate Assessment Plan
                if generate_ap:
                    try:
                        with st.spinner('Generating Assessment Plan and Assessment Summary Record...'):
                            ap_output, asr_output = generate_assessment_documents(context, selected_org)
                        if ap_output:
                            st.success(f"Assessment Plan generated: {ap_output}")
                            st.session_state['ap_output'] = ap_output  # Store output path in session state

                        if asr_output:
                            st.success(f"Assessment Summary Record generated: {asr_output}")
                            st.session_state['asr_output'] = asr_output  # Store output path in session state

                    except Exception as e:
                        st.error(f"Error generating Assessment Documents: {e}")

                # Check if any documents require the timetable
                needs_timetable = (generate_lp or generate_fg)

                # Generate the timetable if needed and not already generated
                if needs_timetable and 'lesson_plan' not in context:
                    try:
                        with st.spinner("Generating Timetable..."):
                            hours = int(''.join(filter(str.isdigit, context["Total_Course_Duration_Hours"])))
                            num_of_days = hours / 8
                            timetable_data = asyncio.run(generate_timetable(context, num_of_days, timetable_openai_struct_model_client))
                            context['lesson_plan'] = timetable_data['lesson_plan']
                        st.session_state['context'] = context  # Update context in session state
                    except Exception as e:
                        st.error(f"Error generating timetable: {e}")
                        return  # Exit if timetable generation fails
                    
                # Now generate Lesson Plan
                if generate_lp:
                    try:
                        with st.spinner("Generating Lesson Plan..."):
                            print("Context for LP:")
                            print(context)
                            # After generating the initial timetable
                            if 'lesson_plan' in context:
                                # Optimize the timetable to ensure proper LU duration utilization
                                context = optimize_timetable(context)
                            print("Optimized Context for LP:")
                            print(context['lesson_plan'])
                            lp_output = generate_lesson_plan(context, selected_org)
                        if lp_output:
                            st.success(f"Lesson Plan generated: {lp_output}")
                            st.session_state['lp_output'] = lp_output  # Store output path in session state
     
                    except Exception as e:
                        st.error(f"Error generating Lesson Plan: {e}")

                # Generate Facilitator's Guide
                if generate_fg:
                    try:
                        with st.spinner("Generating Facilitator's Guide..."):
                            fg_output = generate_facilitators_guide(context, selected_org)
                        if fg_output:
                            st.success(f"Facilitator's Guide generated: {fg_output}")
                            st.session_state['fg_output'] = fg_output  # Store output path in session state

                    except Exception as e:
                        st.error(f"Error generating Facilitator's Guide: {e}")
            else:
                st.error("Context is empty. Cannot proceed with document generation.")
        else:
            st.error("Please upload a CP document and select a Name of Organisation.")

    # Check if any courseware document was generated
    if any([
        st.session_state.get('lg_output'),
        st.session_state.get('ap_output'),
        st.session_state.get('asr_output'),
        st.session_state.get('lp_output'),
        st.session_state.get('fg_output')
    ]):
        st.subheader("Download All Generated Documents as ZIP")

        # Create an in-memory ZIP file
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
            
            # Helper function to add a file to the zip archive
            def add_file(file_path, prefix):
                if file_path and os.path.exists(file_path):
                    # Determine file name based on TGS_Ref_No (if available) or fallback to course title
                    if 'TGS_Ref_No' in st.session_state['context'] and st.session_state['context']['TGS_Ref_No']:
                        file_name = f"{prefix}_{st.session_state['context']['TGS_Ref_No']}_{st.session_state['context']['Course_Title']}_v1.docx"
                    else:
                        file_name = f"{prefix}_{st.session_state['context']['Course_Title']}_v1.docx"
                    zipf.write(file_path, arcname=file_name)
            
            # Add each generated document if it exists
            add_file(st.session_state.get('lg_output'), "LG")
            add_file(st.session_state.get('ap_output'), "Assessment_Plan")
            add_file(st.session_state.get('asr_output'), "Assessment_Summary_Record")
            add_file(st.session_state.get('lp_output'), "LP")
            add_file(st.session_state.get('fg_output'), "FG")
        
        # Reset the buffer's position to the beginning
        zip_buffer.seek(0)
        
        # Create a download button for the ZIP archive
        st.download_button(
            label="Download All Documents (ZIP)",
            data=zip_buffer.getvalue(),
            file_name="courseware_documents.zip",
            mime="application/zip"
        )