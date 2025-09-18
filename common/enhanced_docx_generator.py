"""
Enhanced DOCX Generator with CSS Style Preservation
==================================================

This module provides enhanced HTML-to-DOCX conversion that preserves CSS styling
from the HTML template, mapping CSS properties to equivalent DOCX formatting.

Author: Claude Code Assistant
Date: 3 March 2025
"""

import re
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


class CSSToDocxConverter:
    """Convert CSS styles to DOCX formatting"""
    
    def __init__(self):
        self.color_map = {
            '#333': RGBColor(51, 51, 51),
            '#2b5a87': RGBColor(43, 90, 135),
            '#666': RGBColor(102, 102, 102),
            '#fff': RGBColor(255, 255, 255),
            '#0066cc': RGBColor(0, 102, 204),
            'white': RGBColor(255, 255, 255),
            'black': RGBColor(0, 0, 0),
        }
        
    def parse_css_styles(self, html_content: str) -> Dict[str, Dict[str, str]]:
        """Extract CSS styles from HTML content"""
        css_styles = {}
        
        # Extract style tag content
        soup = BeautifulSoup(html_content, 'html.parser')
        style_tag = soup.find('style')
        
        if not style_tag:
            return css_styles
            
        css_content = style_tag.get_text()
        
        # Parse CSS rules
        css_rules = re.findall(r'([^{]+)\s*\{([^}]+)\}', css_content)
        
        for selector, properties in css_rules:
            selector = selector.strip()
            props = {}
            
            # Parse properties
            for prop in properties.split(';'):
                if ':' in prop:
                    key, value = prop.split(':', 1)
                    props[key.strip()] = value.strip()
            
            css_styles[selector] = props
            
        return css_styles
    
    def get_color_from_hex(self, hex_color: str) -> Optional[RGBColor]:
        """Convert hex color to RGBColor"""
        if hex_color in self.color_map:
            return self.color_map[hex_color]
        
        if hex_color.startswith('#') and len(hex_color) == 7:
            try:
                r = int(hex_color[1:3], 16)
                g = int(hex_color[3:5], 16)
                b = int(hex_color[5:7], 16)
                return RGBColor(r, g, b)
            except ValueError:
                pass
        
        return None
    
    def apply_css_to_run(self, run, css_props: Dict[str, str]):
        """Apply CSS properties to a DOCX run"""
        if 'font-size' in css_props:
            size_str = css_props['font-size'].replace('px', '').replace('pt', '')
            try:
                size = int(size_str)
                run.font.size = Pt(size)
            except ValueError:
                pass
        
        if 'font-weight' in css_props and css_props['font-weight'] == 'bold':
            run.bold = True
        
        if 'font-style' in css_props and css_props['font-style'] == 'italic':
            run.italic = True
        
        if 'color' in css_props:
            color = self.get_color_from_hex(css_props['color'])
            if color:
                run.font.color.rgb = color
        
        if 'text-decoration' in css_props and 'underline' in css_props['text-decoration']:
            run.underline = True
    
    def apply_css_to_paragraph(self, paragraph, css_props: Dict[str, str]):
        """Apply CSS properties to a DOCX paragraph"""
        if 'text-align' in css_props:
            align_map = {
                'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
                'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
                'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
                'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            }
            alignment = align_map.get(css_props['text-align'])
            if alignment:
                paragraph.alignment = alignment
        
        if 'margin-bottom' in css_props:
            margin_str = css_props['margin-bottom'].replace('px', '').replace('pt', '')
            try:
                margin = int(margin_str)
                paragraph.space_after = Pt(margin)
            except ValueError:
                pass
        
        if 'margin-top' in css_props:
            margin_str = css_props['margin-top'].replace('px', '').replace('pt', '')
            try:
                margin = int(margin_str)
                paragraph.space_before = Pt(margin)
            except ValueError:
                pass


class EnhancedDocxGenerator:
    """Enhanced DOCX generator with CSS style preservation"""
    
    def __init__(self):
        self.converter = CSSToDocxConverter()
        self.css_styles = {}
    
    def create_styled_document(self, html_content: str, course_title: str) -> Document:
        """Create a DOCX document with CSS styles applied"""
        
        # Parse CSS styles
        self.css_styles = self.converter.parse_css_styles(html_content)
        
        # Create document
        doc = Document()
        
        # Set document margins (equivalent to container padding)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
        
        # Parse HTML content
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove style and script tags
        for tag in soup(['style', 'script']):
            tag.decompose()
        
        # Process the container div
        container = soup.find('div', class_='container')
        if container:
            self._process_element(doc, container)
        else:
            self._process_element(doc, soup.body or soup)
        
        return doc
    
    def _process_element(self, doc_or_parent, element):
        """Process HTML element and convert to DOCX"""
        
        if not hasattr(element, 'name') or element.name is None:
            # Text node (NavigableString)
            text = str(element).strip()
            if text:
                if hasattr(doc_or_parent, 'add_paragraph'):
                    p = doc_or_parent.add_paragraph()
                    run = p.add_run(text)
                    self._apply_inherited_styles(run, p)
            return
        
        tag_name = element.name
        
        # Handle different HTML tags
        if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self._process_heading(doc_or_parent, element, tag_name)
        
        elif tag_name == 'p':
            self._process_paragraph(doc_or_parent, element)
        
        elif tag_name in ['ul', 'ol']:
            self._process_list(doc_or_parent, element, tag_name)
        
        elif tag_name == 'table':
            self._process_table(doc_or_parent, element)
        
        elif tag_name == 'div':
            self._process_div(doc_or_parent, element)
        
        elif tag_name == 'br':
            if hasattr(doc_or_parent, 'add_paragraph'):
                doc_or_parent.add_paragraph()
        
        elif tag_name in ['strong', 'b', 'em', 'i', 'a', 'span']:
            # Inline elements - these should be handled by parent processing
            pass
        
        else:
            # For other tags, process children
            for child in element.children:
                self._process_element(doc_or_parent, child)
    
    def _process_heading(self, doc_or_parent, element, tag_name):
        """Process heading elements"""
        if hasattr(doc_or_parent, 'add_paragraph'):
            p = doc_or_parent.add_paragraph()
            
            # Apply heading styles based on CSS
            css_props = self.css_styles.get(tag_name, {})
            
            # Process text content with inline formatting
            self._process_inline_content(p, element, css_props)
            
            # Apply paragraph-level styles
            self.converter.apply_css_to_paragraph(p, css_props)
    
    def _process_paragraph(self, doc_or_parent, element):
        """Process paragraph elements"""
        if hasattr(doc_or_parent, 'add_paragraph'):
            p = doc_or_parent.add_paragraph()
            
            # Get CSS styles for paragraph
            css_props = self._get_element_styles(element)
            
            # Process content
            self._process_inline_content(p, element, css_props)
            
            # Apply paragraph styles
            self.converter.apply_css_to_paragraph(p, css_props)
    
    def _process_inline_content(self, paragraph, element, base_styles=None):
        """Process inline content with formatting"""
        base_styles = base_styles or {}
        
        for child in element.children:
            if hasattr(child, 'name') and child.name is not None:
                if child.name in ['strong', 'b']:
                    run = paragraph.add_run(child.get_text())
                    run.bold = True
                    self.converter.apply_css_to_run(run, base_styles)
                
                elif child.name in ['em', 'i']:
                    run = paragraph.add_run(child.get_text())
                    run.italic = True
                    self.converter.apply_css_to_run(run, base_styles)
                
                elif child.name == 'a':
                    run = paragraph.add_run(child.get_text())
                    run.underline = True
                    # Apply link color if available
                    link_color = self.converter.get_color_from_hex('#0066cc')
                    if link_color:
                        run.font.color.rgb = link_color
                    self.converter.apply_css_to_run(run, base_styles)
                
                elif child.name == 'br':
                    paragraph.add_run().add_break()
                
                else:
                    # For other inline elements, get their styles
                    child_styles = self._get_element_styles(child)
                    merged_styles = {**base_styles, **child_styles}
                    
                    run = paragraph.add_run(child.get_text())
                    self.converter.apply_css_to_run(run, merged_styles)
            else:
                # Text node (NavigableString)
                text = str(child).strip()
                if text:
                    run = paragraph.add_run(text)
                    self.converter.apply_css_to_run(run, base_styles)
    
    def _process_list(self, doc_or_parent, element, list_type):
        """Process list elements"""
        for li in element.find_all('li', recursive=False):
            if hasattr(doc_or_parent, 'add_paragraph'):
                p = doc_or_parent.add_paragraph()
                
                # Add bullet or number
                if list_type == 'ul':
                    bullet_run = p.add_run('• ')
                else:
                    # For ordered lists, we'd need to track numbering
                    bullet_run = p.add_run('1. ')
                
                # Process list item content
                self._process_inline_content(p, li)
    
    def _process_table(self, doc_or_parent, element):
        """Process table elements"""
        if not hasattr(doc_or_parent, 'add_table'):
            return
        
        rows = element.find_all('tr')
        if not rows:
            return
        
        # Determine table dimensions
        max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
        
        # Create table
        table = doc_or_parent.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        
        # Populate table
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            docx_row = table.rows[row_idx]
            
            for col_idx, cell in enumerate(cells):
                if col_idx < len(docx_row.cells):
                    docx_cell = docx_row.cells[col_idx]
                    
                    # Clear default paragraph
                    docx_cell._tc.clear_content()
                    
                    # Add content
                    p = docx_cell.add_paragraph()
                    self._process_inline_content(p, cell)
                    
                    # Apply header styling if th
                    if cell.name == 'th':
                        for paragraph in docx_cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
    
    def _process_div(self, doc_or_parent, element):
        """Process div elements based on their classes"""
        classes = element.get('class', [])
        
        if 'page-break' in classes:
            # Add page break
            if hasattr(doc_or_parent, 'add_page_break'):
                doc_or_parent.add_page_break()
            elif hasattr(doc_or_parent, 'add_paragraph'):
                p = doc_or_parent.add_paragraph()
                p.add_run().add_break(WD_BREAK.PAGE)
        
        # Get styles for this div
        div_styles = self._get_element_styles(element)
        
        # For special sections, add background color simulation
        if any(cls in classes for cls in ['funding-section', 'certificate-section', 'enquiry-section']):
            # Add a paragraph with special formatting to simulate background
            if hasattr(doc_or_parent, 'add_paragraph'):
                p = doc_or_parent.add_paragraph()
                run = p.add_run(" ")  # Empty space to create background effect
                # Note: DOCX doesn't support paragraph background colors easily
                # This would need more advanced formatting
        
        # Process children
        for child in element.children:
            self._process_element(doc_or_parent, child)
    
    def _get_element_styles(self, element) -> Dict[str, str]:
        """Get CSS styles for an element"""
        styles = {}
        
        # Skip if not an actual element (e.g., text node)
        if not hasattr(element, 'name') or element.name is None:
            return styles
        
        # Check for tag-specific styles
        if element.name in self.css_styles:
            styles.update(self.css_styles[element.name])
        
        # Check for class-specific styles
        classes = element.get('class', [])
        for cls in classes:
            class_selector = f'.{cls}'
            if class_selector in self.css_styles:
                styles.update(self.css_styles[class_selector])
        
        # Check for ID-specific styles
        element_id = element.get('id')
        if element_id:
            id_selector = f'#{element_id}'
            if id_selector in self.css_styles:
                styles.update(self.css_styles[id_selector])
        
        return styles
    
    def _apply_inherited_styles(self, run, paragraph):
        """Apply inherited styles to run and paragraph"""
        # Apply body styles as defaults
        body_styles = self.css_styles.get('body', {})
        self.converter.apply_css_to_run(run, body_styles)
        self.converter.apply_css_to_paragraph(paragraph, body_styles)


def generate_enhanced_docx(html_content: str, course_title: str, output_path: str) -> bool:
    """
    Generate enhanced DOCX with CSS style preservation
    
    Args:
        html_content (str): HTML content with CSS styles
        course_title (str): Course title
        output_path (str): Output path for DOCX file
        
    Returns:
        bool: Success status
    """
    try:
        generator = EnhancedDocxGenerator()
        doc = generator.create_styled_document(html_content, course_title)
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error generating enhanced DOCX: {e}")
        return False


if __name__ == "__main__":
    # Test the converter
    template_path = Path("common/brochure_template/brochure.html")
    if template_path.exists():
        with open(template_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        output_path = "test_output.docx"
        if generate_enhanced_docx(html_content, "Test Course", output_path):
            print(f"Enhanced DOCX generated: {output_path}")
        else:
            print("Failed to generate enhanced DOCX")