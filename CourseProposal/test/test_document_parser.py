import sys
import os
import json
import re
from docx import Document

# Copy the relevant functions from document_parser.py 
def parse_document(file_path):
    """
    Parse a Word document and extract structured data.
    
    Args:
        file_path (str): Path to the Word document.
        
    Returns:
        dict: Extracted data organized by sections.
    """
    doc = Document(file_path)
    
    # Initialize data structure
    data = {
        "Course Information": {},
        "Learning Outcomes": {
            "Learning Outcomes": [],
            "Knowledge": [],
            "Ability": [],
            "Knowledge and Ability Mapping": {},
            "Course Duration": "",
            "Instructional Methods": "",
            "content": []
        },
        "TSC and Topics": {
            "TSC Title": "",
            "TSC Code": "",
            "Topics": [],
            "Learning Units": [],
            "content": []
        },
        "Assessment Methods": {
            "Assessment Methods": [],
            "Amount of Practice Hours": "",
            "Course Outline": {
                "Learning Units": {}
            },
            "content": [],
            "Assessment Details": {}
        }
    }
    
    # Parse paragraphs
    current_section = None
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Determine section based on content
        if "Course Title" in text or "Name of Organisation" in text:
            current_section = "Course Information"
            parse_course_info(text, data[current_section])
        elif "Learning Outcome" in text or "LO" in text:
            current_section = "Learning Outcomes"
            data[current_section]["content"].append(text)
        elif "TSC" in text or "Knowledge:" in text or "Abilities:" in text:
            current_section = "TSC and Topics"
            data[current_section]["content"].append(text)
        elif "Assessment Method" in text:
            current_section = "Assessment Methods"
            parse_assessment_methods(text, data[current_section])
        elif current_section:
            # Add content to the current section
            data[current_section]["content"].append(text)
    
    # Parse tables
    for table in doc.tables:
        table_data = parse_table(table)
        
        # Determine which section the table belongs to
        if any(col.startswith("LU") for col in table_data[0] if isinstance(col, str)):
            data["TSC and Topics"]["content"].append({"table": table_data})
    
    return data

def parse_course_info(text, data):
    """Parse course information from text."""
    if "Course Title:" in text:
        data["Course Title"] = text.split("Course Title:")[1].strip()
    elif "Name of Organisation:" in text:
        data["Name of Organisation"] = text.split("Name of Organisation:")[1].strip()

def parse_assessment_methods(text, data):
    """Parse assessment methods from text."""
    if "Assessment Methods:" in text:
        methods_text = text.split("Assessment Methods:")[1].strip()
        methods = [method.strip() for method in methods_text.split(",")]
        data["Assessment Methods"] = methods
        
        # Extract assessment hours from brackets
        assessment_details = {}
        for method in methods:
            match = re.search(r"(.*?)\s*\((.+?)\)", method)
            if match:
                method_name = match.group(1).strip()
                hours = match.group(2).strip()
                assessment_details[method_name] = hours
        
        if assessment_details:
            data["Assessment Details"] = assessment_details

# Function to parse tables (ENHANCED for subtopics)
def parse_table(table):
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = cell.text.strip()
            # Check if this is a topic cell (starts with T# or Topic #)
            if re.match(r"^(T\d+:|Topic \d+:)", cell_text):
                # Split by newlines, first line is topic, rest are subtopics
                lines = [l.strip() for l in cell_text.split('\n') if l.strip()]
                if len(lines) > 1:
                    topic = lines[0]
                    subtopics = [re.sub(r"^[•−–●◦]\s*", "", l) for l in lines[1:]]
                    cell_value = {"Topic": topic, "Details": subtopics}
                    cells.append(cell_value)
                else:
                    cell_value = {"Topic": cell_text, "Details": []}
                    cells.append(cell_value)
            else:
                cells.append(cell_text)
        rows.append(cells)
    return rows

def test_document_parser():
    """Test the document parser's ability to extract subtopics from topic cells."""
    # Input document path
    input_docx = os.path.join('input', 'TSC_Mapping-TI-ai-storytelling_enhanced.docx')
    
    # Parse the document
    print(f"Parsing document: {input_docx}")
    parsed_data = parse_document(input_docx)
    
    # Save the output to a temporary file for inspection
    temp_output_json = os.path.join('json_output', 'test_parser_output.json')
    with open(temp_output_json, 'w', encoding='utf-8') as f:
        json.dump(parsed_data, f, indent=4)
    
    print(f"Parsed output saved to: {temp_output_json}")
    
    # Check for topic cells with subtopics (specifically looking at Topic 1 in LU1)
    print("\nSearching for topic cells with subtopics...")
    for section_name, section_content in parsed_data.items():
        if section_name == "TSC and Topics" and "content" in section_content:
            for item in section_content["content"]:
                if isinstance(item, dict) and "table" in item:
                    table = item["table"]
                    print("\nFound table in TSC and Topics section")
                    
                    # Check each row in the table
                    for row_idx, row in enumerate(table):
                        if row_idx == 0:  # Skip header row
                            continue
                        
                        # Check if this row has a topic cell (usually in position 3)
                        if len(row) > 3 and isinstance(row[3], dict) and "Topic" in row[3]:
                            topic = row[3]["Topic"]
                            details = row[3].get("Details", [])
                            print(f"\nRow {row_idx}:")
                            print(f"  Topic: {topic}")
                            print(f"  Details: {details}")
                            
                            # Special check for Topic 1 in LU1
                            if "T1: Fundamentals of storytelling" in topic:
                                print(f"  *** FOUND Topic 1 with {len(details)} subtopics ***")
                                for idx, detail in enumerate(details):
                                    print(f"    Subtopic {idx+1}: {detail}")

if __name__ == "__main__":
    test_document_parser() 