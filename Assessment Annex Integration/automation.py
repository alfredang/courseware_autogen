import os
import io
import re
import streamlit as st
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google.oauth2 import service_account
from docx import Document
from docxcompose.composer import Composer

# -------------------------------------------------------------------
# 0. AUTHENTICATION
# -------------------------------------------------------------------

def authenticate():
    creds = None
    try:
        creds = service_account.Credentials.from_service_account_info(
            st.secrets["GOOGLE_API_CREDS"]
        )
        return creds
    except Exception as e:
        print(f"Error during authentication: {e}")
        return None

# -------------------------------------------------------------------
# 1. HELPER FUNCTIONS
# -------------------------------------------------------------------

def get_version_tuple(filename):
    """Extract version number from filename, e.g., '_v1.2' -> (1,2)."""
    pattern = r'(?:-|_)v(\d+(\.\d+)*)'
    match = re.search(pattern, filename, re.IGNORECASE)
    if match:
        return tuple(map(int, match.group(1).split('.')))
    return (0,)

def is_question_file(filename, abbreviations):
    """Check if filename is a valid question docx file."""
    abbr_pattern = '|'.join(map(re.escape, abbreviations))
    return re.match(fr'^({abbr_pattern}).*\.docx$', filename, re.IGNORECASE)

def is_answer_file(filename, abbreviations):
    """Check if filename is a valid answer docx file."""
    abbr_pattern = '|'.join(map(re.escape, abbreviations))
    return re.match(fr'^(Answer to |Answers to )({abbr_pattern}).*\.docx$', filename, re.IGNORECASE)

def download_docx(file_id, file_name, drive_service, download_dir="./downloads"):
    """Download a Google Drive file and save it as .docx."""
    if not os.path.exists(download_dir):
        os.makedirs(download_dir, exist_ok=True)

    request = drive_service.files().get_media(fileId=file_id)
    file_path = os.path.join(download_dir, file_name)
    with io.BytesIO() as fh:
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()
        fh.seek(0)
        with open(file_path, "wb") as f:
            f.write(fh.read())
    return file_path

# -------------------------------------------------------------------
# 2. PROCESS COURSE FOLDER
# -------------------------------------------------------------------

def process_course_folder(course_folder_id, drive_service, abbreviations):
    """
    Process a course folder to find the latest Assessment Plan and Q&A files.
    :param course_folder_id: ID of the course folder.
    :param drive_service: Authenticated Google Drive service.
    :param abbreviations: List of assessment method abbreviations (e.g., WA (SAQ), PP).
    :return: Dictionary with Assessment Plan and Q&A files.
    """
    # Get subfolders
    subfolders = drive_service.files().list(
        q=f"'{course_folder_id}' in parents and mimeType='application/vnd.google-apps.folder'"
    ).execute().get('files', [])

    assessment_plan_folder = next((f for f in subfolders if f['name'] == 'Assessment Plan'), None)
    assessment_folder = next((f for f in subfolders if f['name'] == 'Assessment'), None)

    if not assessment_plan_folder or not assessment_folder:
        return None

    # Find the latest Assessment Plan
    plan_files = drive_service.files().list(
        q=f"'{assessment_plan_folder['id']}' in parents and "
          "mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'",
    ).execute().get('files', [])

    if not plan_files:
        return None

    latest_plan = max(plan_files, key=lambda f: get_version_tuple(f['name']))

    # Find Q&A files
    assess_files = drive_service.files().list(
        q=f"'{assessment_folder['id']}' in parents and "
          "mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'",
    ).execute().get('files', [])

    method_data = {}
    for abbr in abbreviations:
        q_file = next((f for f in assess_files if is_question_file(f['name'], [abbr])), None)
        a_file = next((f for f in assess_files if is_answer_file(f['name'], [abbr])), None)
        method_data[abbr] = {'question': q_file, 'answer': a_file}

    return {'assessment_plan': latest_plan, 'method_data': method_data}

# -------------------------------------------------------------------
# 3. MERGE DOCUMENTS WITH PLACEHOLDER REPLACEMENT
# -------------------------------------------------------------------

def insert_docs_under_heading(plan_path, heading_map, method_data):
    """
    Replace placeholders in the Assessment Plan with Q&A content.
    :param plan_path: Path to the Assessment Plan.
    :param heading_map: Mapping of placeholders to abbreviations.
    :param method_data: Dictionary of question and answer file paths.
    :return: Path to the updated document.
    """
    base_doc = Document(plan_path)
    composer = Composer(base_doc)

    for placeholder, abbr in heading_map.items():
        if abbr in method_data:
            files = method_data[abbr]
            q_file = files.get('question')
            a_file = files.get('answer')

            # Append question content
            if q_file and 'local_path' in q_file:
                composer.append(Document(q_file['local_path']))

            # Append answer content
            if a_file and 'local_path' in a_file:
                composer.append(Document(a_file['local_path']))

    updated_path = plan_path.replace(".docx", "_Updated.docx")
    composer.save(updated_path)
    return updated_path

# -------------------------------------------------------------------
# 4. MAIN FUNCTION
# -------------------------------------------------------------------

def main():
    creds = authenticate()
    if not creds:
        return

    drive_service = build('drive', 'v3', credentials=creds)

    # Top-level folder name
    top_folder_name = "Assessment Automation"
    top_folder = drive_service.files().list(
        q=f"name='{top_folder_name}' and mimeType='application/vnd.google-apps.folder'"
    ).execute().get('files', [])

    if not top_folder:
        print("Top-level folder not found.")
        return

    top_folder_id = top_folder[0]['id']
    course_folders = drive_service.files().list(
        q=f"'{top_folder_id}' in parents and mimeType='application/vnd.google-apps.folder'"
    ).execute().get('files', [])

    # Abbreviations and heading map
    abbreviations = ["WA (SAQ)", "PP", "CS", "RP", "Oral Questioning"]
    heading_map = {
        "Assessment Questions and Answers for WA(SAQ)": "WA (SAQ)",
        "Assessment Questions and Practical Performance": "PP",
        "Assessment Questions and Oral Questioning (OQ)": "Oral Questioning",
        "Assessment Questions and Case Study": "CS"
    }

    for course_folder in course_folders:
        print(f"\nProcessing Course Folder: {course_folder['name']}")
        result = process_course_folder(course_folder['id'], drive_service, abbreviations)

        if not result:
            print(f"Skipping {course_folder['name']} due to missing files.")
            continue

        assessment_plan = result['assessment_plan']
        method_data = result['method_data']

        # Download Assessment Plan
        plan_path = download_docx(assessment_plan['id'], assessment_plan['name'], drive_service)

        # Download Q&A files
        for abbr, files in method_data.items():
            for doc_type in ['question', 'answer']:
                if files[doc_type]:
                    local_path = download_docx(
                        files[doc_type]['id'], files[doc_type]['name'], drive_service
                    )
                    files[doc_type]['local_path'] = local_path

        # Merge documents
        updated_doc_path = insert_docs_under_heading(plan_path, heading_map, method_data)
        print(f"Merged document saved at: {updated_doc_path}")

if __name__ == "__main__":
    main()
