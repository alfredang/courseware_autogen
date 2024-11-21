from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# Function to set the font to Arial and the font size
def set_font_to_arial(paragraph):
    for run in paragraph.runs:
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')  # Ensure compatibility for non-ASCII characters if needed

# Function to set line spacing to 1.0
def set_line_spacing(paragraph):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = Pt(12)  # This sets line spacing to single (1.0)

# Load the document
doc = Document('Updated_CP_Document.docx')  # Replace with your .docx file path

# Iterate over all paragraphs in the document
for para in doc.paragraphs:
    set_font_to_arial(para)
    set_line_spacing(para)

# Iterate over all tables in the document, because tables may have text that needs formatting as well
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                set_font_to_arial(para)
                set_line_spacing(para)

# Save the updated document
doc.save('Updated_CP_Document2.docx')  # Replace with your desired output file name
